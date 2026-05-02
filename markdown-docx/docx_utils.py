#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from copy import deepcopy
from dataclasses import dataclass
from typing import Iterable, List, Optional

from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class RunStyle:
    ascii_font: str = "Times New Roman"
    east_asia_font: str = "宋体"
    size_pt: float = 12.0
    bold: Optional[bool] = None


CITATION_PATTERN = re.compile(r"(\[(?:\d+\s*(?:[-,，、]\s*\d+\s*)*)\])")


def iter_block_items(parent: DocumentObject) -> Iterable[object]:
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def clear_paragraph(paragraph: Paragraph) -> None:
    node = paragraph._p
    for child in list(node):
        if child.tag.endswith("pPr"):
            continue
        node.remove(child)


def remove_block(block: object) -> None:
    element = block._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def has_section_break(paragraph: Paragraph) -> bool:
    try:
        ppr = paragraph._p.pPr
        return bool(ppr is not None and ppr.sectPr is not None)
    except Exception:
        return False


def remove_or_clear_block(block: object) -> None:
    if isinstance(block, Paragraph) and has_section_break(block):
        clear_paragraph(block)
        return
    remove_block(block)


def copy_paragraph_properties(target: Paragraph, prototype: Paragraph) -> None:
    target_node = target._p
    for child in list(target_node):
        if child.tag.endswith("pPr"):
            target_node.remove(child)
    if prototype._p.pPr is not None:
        target_node.insert(0, deepcopy(prototype._p.pPr))
    if prototype.style is not None:
        target.style = prototype.style


def insert_paragraph_before(reference: Paragraph, prototype: Optional[Paragraph] = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    reference._p.addprevious(new_p)
    paragraph = Paragraph(new_p, reference._parent)
    if prototype is not None:
        copy_paragraph_properties(paragraph, prototype)
    return paragraph


def first_nonempty_run_style(paragraph: Paragraph, fallback_font: str = "宋体", fallback_size: float = 12.0) -> RunStyle:
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        rfonts = None
        try:
            rfonts = run._element.rPr.rFonts
        except Exception:
            rfonts = None
        east_asia_font = fallback_font
        ascii_font = run.font.name or "Times New Roman"
        if rfonts is not None:
            east_asia_font = rfonts.get(qn("w:eastAsia")) or run.font.name or fallback_font
            ascii_font = rfonts.get(qn("w:ascii")) or run.font.name or "Times New Roman"
        elif run.font.name:
            east_asia_font = run.font.name
        size_pt = run.font.size.pt if run.font.size else fallback_size
        return RunStyle(
            ascii_font=ascii_font,
            east_asia_font=east_asia_font,
            size_pt=size_pt,
            bold=run.bold,
        )
    return RunStyle(east_asia_font=fallback_font, size_pt=fallback_size)


def ensure_rfonts(run, ascii_font: str, east_asia_font: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def _apply_run_style(run, style: RunStyle, east_asia_font: Optional[str] = None, superscript: bool = False) -> None:
    from docx.shared import Pt

    run.font.name = style.ascii_font
    ensure_rfonts(
        run,
        ascii_font=style.ascii_font,
        east_asia_font=east_asia_font or style.east_asia_font,
    )
    if style.bold is not None:
        run.bold = style.bold
    if style.size_pt:
        run.font.size = Pt(style.size_pt)
    run.font.superscript = superscript


def add_styled_run(paragraph: Paragraph, text: str, style: RunStyle, superscript: bool = False) -> None:
    if not text:
        return

    buffer: List[str] = []
    current_is_cjk: Optional[bool] = None

    def flush(superscript_chunk: bool = False) -> None:
        nonlocal buffer, current_is_cjk
        if not buffer:
            return
        chunk = "".join(buffer)
        run = paragraph.add_run(chunk)
        east_asia_font = style.east_asia_font if current_is_cjk else style.ascii_font
        _apply_run_style(run, style, east_asia_font=east_asia_font, superscript=superscript_chunk)
        buffer = []
        current_is_cjk = None

    if superscript:
        buffer.append(text)
        flush(superscript_chunk=True)
        return

    for char in text:
        is_cjk = bool(re.match(r"[\u4e00-\u9fff]", char))
        if current_is_cjk is None:
            current_is_cjk = is_cjk
            buffer.append(char)
            continue
        if is_cjk == current_is_cjk:
            buffer.append(char)
            continue
        flush()
        current_is_cjk = is_cjk
        buffer.append(char)
    flush()


def add_mixed_text(paragraph: Paragraph, text: str, style: RunStyle, superscript_citations: bool = False) -> None:
    if not text:
        return

    if not superscript_citations:
        add_styled_run(paragraph, text, style)
        return

    parts = CITATION_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if CITATION_PATTERN.fullmatch(part):
            add_styled_run(paragraph, part, style, superscript=True)
        else:
            add_styled_run(paragraph, part, style)


def set_paragraph_text(paragraph: Paragraph, text: str, style: RunStyle, superscript_citations: bool = False) -> None:
    clear_paragraph(paragraph)
    add_mixed_text(paragraph, text, style, superscript_citations=superscript_citations)


def set_cell_text(cell, text: str, style: RunStyle) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text, style)


def set_keyword_paragraph(paragraph: Paragraph, prefix: str, value: str, prefix_style: RunStyle, value_style: RunStyle) -> None:
    clear_paragraph(paragraph)
    prefix_run = paragraph.add_run(prefix)
    from docx.shared import Pt

    prefix_run.font.name = prefix_style.ascii_font
    prefix_run.font.size = Pt(prefix_style.size_pt)
    prefix_run.bold = True if prefix_style.bold is None else prefix_style.bold
    ensure_rfonts(prefix_run, prefix_style.ascii_font, prefix_style.east_asia_font)
    add_mixed_text(paragraph, value, value_style)


def insert_table_before(doc: DocumentObject, reference: Paragraph, rows: int, cols: int, style: Optional[str] = None) -> Table:
    table = doc.add_table(rows=rows, cols=cols)
    if style:
        try:
            table.style = style
        except Exception:
            pass
    reference._p.addprevious(table._tbl)
    return table


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def remove_between(first_block: object, end_paragraph: Paragraph) -> None:
    current = first_block
    while current is not None and current._element is not end_paragraph._element:
        next_element = current._element.getnext()
        remove_block(current)
        if next_element is None:
            current = None
            continue
        if isinstance(next_element, CT_P):
            current = Paragraph(next_element, end_paragraph._parent)
        elif isinstance(next_element, CT_Tbl):
            current = Table(next_element, end_paragraph._parent)
        else:
            current = None


def body_blocks_between(doc: DocumentObject, start_paragraph: Paragraph, end_paragraph: Paragraph) -> List[object]:
    blocks: List[object] = []
    capture = False
    for block in iter_block_items(doc):
        if block._element is start_paragraph._element:
            capture = True
            continue
        if block._element is end_paragraph._element:
            break
        if capture:
            blocks.append(block)
    return blocks


def first_paragraph_between(doc: DocumentObject, start_paragraph: Paragraph, end_paragraph: Paragraph) -> Optional[Paragraph]:
    for block in body_blocks_between(doc, start_paragraph, end_paragraph):
        if isinstance(block, Paragraph):
            return block
    return None
