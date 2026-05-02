#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
from pathlib import Path

from docx import Document

from docx_utils import clear_paragraph, first_paragraph_between
from header_utils import update_headers
from legacy_format_rules import (
    BODY_STYLE,
    ENGLISH_BODY_STYLE,
    ENGLISH_KEYWORD_LABEL_STYLE,
    KEYWORD_LABEL_STYLE,
    MAJOR_HEADING_STYLE,
    SUBHEADING_STYLE,
    format_body,
    format_keywords,
    format_major_heading,
    format_major_heading_on_new_page,
    format_reference,
    format_subheading,
)
from template_locator import find_next_anchor, load_schema, locate_anchors
from docx_utils import set_keyword_paragraph, set_paragraph_text


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_SCHEMA = SCRIPT_DIR / "dhu_template_schema.json"
DEFAULT_TEMPLATE = SCRIPT_DIR / "outputs" / "dhu_undergrad_clean_template.docx"


def optimize_clean_template(template_path: Path, schema_path: Path) -> None:
    schema = load_schema(schema_path)
    placeholders = schema["clean_placeholders"]
    doc = Document(str(template_path))
    anchors = locate_anchors(doc, schema, required=True)

    update_headers(doc, placeholders["front_cn_title"])

    body_start = doc.paragraphs[anchors["body_start"]]
    body_h2 = doc.paragraphs[anchors["body_start"] + 1]
    body_h3 = doc.paragraphs[anchors["body_start"] + 2]
    body_text = doc.paragraphs[anchors["body_start"] + 3]

    set_paragraph_text(body_start, body_start.text, MAJOR_HEADING_STYLE)
    format_major_heading(body_start)

    set_paragraph_text(body_h2, body_h2.text, SUBHEADING_STYLE)
    format_subheading(body_h2)

    set_paragraph_text(body_h3, body_h3.text, SUBHEADING_STYLE)
    format_subheading(body_h3)

    if body_text.text.strip():
        set_paragraph_text(body_text, body_text.text, BODY_STYLE)
    format_body(body_text)
    if not body_text.text.strip():
        clear_paragraph(body_text)

    cn_keywords = doc.paragraphs[anchors["cn_keywords"]]
    set_keyword_paragraph(cn_keywords, "关键词：", "", KEYWORD_LABEL_STYLE, BODY_STYLE)
    format_keywords(cn_keywords)

    en_keywords = doc.paragraphs[anchors["en_keywords"]]
    set_keyword_paragraph(en_keywords, "KEY WORDS: ", "", ENGLISH_KEYWORD_LABEL_STYLE, ENGLISH_BODY_STYLE)
    format_keywords(en_keywords)

    reference_end_name = find_next_anchor(anchors, "reference", "appendix", "acknowledgements", "foreign_translation")
    reference_end = doc.paragraphs[anchors[reference_end_name]] if reference_end_name else doc.paragraphs[-1]
    format_major_heading_on_new_page(doc.paragraphs[anchors["reference"]])
    reference_slot = first_paragraph_between(doc, doc.paragraphs[anchors["reference"]], reference_end)
    if reference_slot is not None:
        format_reference(reference_slot)
        if not reference_slot.text.strip():
            clear_paragraph(reference_slot)

    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="按旧 formatter 规则校正现有干净模板的占位格式")
    parser.add_argument(
        "-t",
        "--template",
        default=str(DEFAULT_TEMPLATE),
        help="待优化的干净模板路径",
    )
    parser.add_argument(
        "-s",
        "--schema",
        default=str(DEFAULT_SCHEMA),
        help="模板 schema 路径",
    )
    args = parser.parse_args()

    optimize_clean_template(Path(args.template).resolve(), Path(args.schema).resolve())
    print(f"已优化干净模板: {Path(args.template).resolve()}")


if __name__ == "__main__":
    main()
