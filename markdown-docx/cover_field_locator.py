#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from docx_utils import RunStyle, first_nonempty_run_style, iter_block_items, set_paragraph_text


def _normalize_label(text: str) -> str:
    cleaned = re.sub(r"[：:\s]+", "", (text or "").strip())
    return cleaned


def locate_cover_cells(doc: Document, field_map: Dict[str, str]) -> Dict[str, _Cell]:
    normalized = {key: _normalize_label(value) for key, value in field_map.items()}
    result: Dict[str, _Cell] = {}
    for table in doc.tables[:3]:
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                cell_text = _normalize_label(cell.text)
                for key, label in normalized.items():
                    if cell_text != label:
                        continue
                    if index + 1 < len(row.cells):
                        result[key] = row.cells[index + 1]
    return result


def _nonempty_paragraphs_before_first_table(doc: Document) -> List[Paragraph]:
    paragraphs: List[Paragraph] = []
    for block in iter_block_items(doc):
        if isinstance(block, Table):
            break
        if isinstance(block, Paragraph) and block.text.strip():
            paragraphs.append(block)
    return paragraphs


def locate_cover_title_paragraphs(doc: Document) -> Tuple[Optional[Paragraph], Optional[Paragraph]]:
    paragraphs = _nonempty_paragraphs_before_first_table(doc)
    if len(paragraphs) < 2:
        return None, None
    return paragraphs[-2], paragraphs[-1]


def locate_front_title_paragraph(doc: Document, cn_abstract_idx: int) -> Optional[Paragraph]:
    for idx in range(cn_abstract_idx - 1, -1, -1):
        paragraph = doc.paragraphs[idx]
        if paragraph.text.strip():
            return paragraph
    return None


def locate_front_english_title_paragraphs(doc: Document, cn_keywords_idx: int, en_abstract_idx: int) -> List[Paragraph]:
    paragraphs: List[Paragraph] = []
    for idx in range(cn_keywords_idx + 1, en_abstract_idx):
        paragraph = doc.paragraphs[idx]
        if paragraph.text.strip():
            paragraphs.append(paragraph)
    return paragraphs


def set_cell_text(cell: _Cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    style = first_nonempty_run_style(paragraph)
    set_paragraph_text(paragraph, text, style)


def fill_cover_fields(doc: Document, schema: Dict, metadata: Dict[str, str], blank_missing: bool = True) -> None:
    cell_map = locate_cover_cells(doc, schema["cover_fields"])
    for key, cell in cell_map.items():
        value = metadata.get(key, "")
        if not value and not blank_missing:
            continue
        set_cell_text(cell, value)


def set_cover_titles(
    doc: Document,
    cn_abstract_idx: int,
    cn_keywords_idx: int,
    en_abstract_idx: int,
    chinese_title: str,
    english_title_lines: List[str],
) -> None:
    cover_cn, cover_en = locate_cover_title_paragraphs(doc)
    if cover_cn is not None:
        cn_style = first_nonempty_run_style(cover_cn, fallback_font="黑体", fallback_size=22.0)
        cn_style.bold = True if cn_style.bold is None else cn_style.bold
        set_paragraph_text(cover_cn, chinese_title, cn_style)
    if cover_en is not None:
        en_style = first_nonempty_run_style(cover_en, fallback_font="Times New Roman", fallback_size=22.0)
        en_style.bold = True if en_style.bold is None else en_style.bold
        set_paragraph_text(cover_en, english_title_lines[0] if english_title_lines else "", en_style)

    front_cn = locate_front_title_paragraph(doc, cn_abstract_idx)
    if front_cn is not None:
        front_cn_style = first_nonempty_run_style(front_cn, fallback_font="黑体", fallback_size=16.0)
        front_cn_style.bold = True if front_cn_style.bold is None else front_cn_style.bold
        set_paragraph_text(front_cn, chinese_title, front_cn_style)

    front_en_paragraphs = locate_front_english_title_paragraphs(doc, cn_keywords_idx, en_abstract_idx)
    if not front_en_paragraphs:
        return
    while len(english_title_lines) < len(front_en_paragraphs):
        english_title_lines.append("")
    for paragraph, line in zip(front_en_paragraphs, english_title_lines):
        style = first_nonempty_run_style(paragraph, fallback_font="Times New Roman", fallback_size=16.0)
        style.bold = True if style.bold is None else style.bold
        set_paragraph_text(paragraph, line, style)


def wrap_english_title(title: str, max_lines: int = 3, approx_line_length: int = 34) -> List[str]:
    words = title.split()
    if not words:
        return [""] * max_lines

    lines: List[str] = []
    current: List[str] = []
    current_length = 0
    for word in words:
        projected = current_length + (1 if current else 0) + len(word)
        if current and projected > approx_line_length and len(lines) < max_lines - 1:
            lines.append(" ".join(current))
            current = [word]
            current_length = len(word)
        else:
            current.append(word)
            current_length = projected
    if current:
        lines.append(" ".join(current))
    while len(lines) < max_lines:
        lines.append("")
    return lines[:max_lines]
