#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import re
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Cm, Emu

from cover_field_locator import fill_cover_fields, set_cover_titles, wrap_english_title
from docx_utils import (
    RunStyle,
    body_blocks_between,
    clear_paragraph,
    copy_paragraph_properties,
    first_nonempty_run_style,
    first_paragraph_between,
    has_section_break,
    insert_paragraph_before,
    insert_table_before,
    remove_or_clear_block,
    set_cell_text,
    set_keyword_paragraph,
    set_paragraph_text,
    set_repeat_table_header,
)
from header_utils import update_headers
from legacy_format_rules import (
    BODY_STYLE,
    CONTINUED_TABLE_LABEL_STYLE,
    ENGLISH_BODY_STYLE,
    ENGLISH_KEYWORD_LABEL_STYLE,
    FIGURE_CAPTION_STYLE,
    KEYWORD_LABEL_STYLE,
    MAJOR_HEADING_STYLE,
    SUBHEADING_STYLE,
    TABLE_CAPTION_STYLE,
    TABLE_CELL_STYLE,
    format_body,
    format_continued_table_label,
    format_figure_caption,
    format_keywords,
    format_major_heading,
    format_major_heading_on_new_page,
    format_picture_block,
    format_reference,
    format_subheading,
    format_table_caption,
)
from markdown_parser import merge_content_with_metadata, parse_markdown, load_metadata
from template_locator import find_next_anchor, load_schema, locate_anchors


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_SCHEMA = SCRIPT_DIR / "dhu_template_schema.json"
DEFAULT_TEMPLATE = SCRIPT_DIR / "outputs" / "dhu_undergrad_clean_template.docx"
ENABLE_HEURISTIC_CONTINUED_TABLES = False
FIRST_TABLE_PAGE_CAPACITY = 34
CONTINUED_TABLE_PAGE_CAPACITY = 38


def _prepare_slot(doc: Document, start_paragraph, end_paragraph, preserve_start: bool):
    if preserve_start:
        for block in body_blocks_between(doc, start_paragraph, end_paragraph):
            remove_or_clear_block(block)
        clear_paragraph(start_paragraph)
        return start_paragraph

    paragraph = first_paragraph_between(doc, start_paragraph, end_paragraph)
    if paragraph is None:
        paragraph = insert_paragraph_before(end_paragraph)
        copy_paragraph_properties(paragraph, start_paragraph)
    else:
        for block in body_blocks_between(doc, paragraph, end_paragraph):
            remove_or_clear_block(block)
        clear_paragraph(paragraph)
    return paragraph


def _write_text_blocks(slot_paragraph, paragraphs: List[str], style: RunStyle, end_paragraph, formatter=None) -> None:
    if not paragraphs:
        clear_paragraph(slot_paragraph)
        return
    set_paragraph_text(slot_paragraph, paragraphs[0], style)
    if formatter is not None:
        formatter(slot_paragraph)
    for text in paragraphs[1:]:
        new_paragraph = insert_paragraph_before(end_paragraph, slot_paragraph)
        set_paragraph_text(new_paragraph, text, style)
        if formatter is not None:
            formatter(new_paragraph)


def _write_references(doc: Document, slot_paragraph, references: List[str], end_paragraph) -> None:
    if not references:
        clear_paragraph(slot_paragraph)
        return

    insertion_boundary = end_paragraph
    for block in body_blocks_between(doc, slot_paragraph, end_paragraph):
        if getattr(block, "_element", None) is None:
            continue
        if block.__class__.__name__ == "Paragraph" and has_section_break(block):
            insertion_boundary = block
            break

    for number, raw_text in enumerate(references, 1):
        cleaned = re.sub(r"^\[\d+\]\s*", "", raw_text.strip())
        cleaned = re.sub(r"^\d+\.\s*", "", cleaned)
        paragraph = slot_paragraph if number == 1 else insert_paragraph_before(insertion_boundary, slot_paragraph)
        set_paragraph_text(paragraph, f"[{number}] {cleaned}", BODY_STYLE)
        format_reference(paragraph)


def _normalize_caption_text(text: str, fallback: str) -> str:
    cleaned = (text or "").strip()
    return cleaned or fallback


def _body_section_for_media(doc: Document):
    if len(doc.sections) >= 3:
        return doc.sections[2]
    return doc.sections[-1]


def _available_body_width(section) -> Emu:
    return Emu(int(section.page_width) - int(section.left_margin) - int(section.right_margin))


def _resolve_media_path(media_path: str, markdown_dir: Path) -> Path:
    candidate = Path(media_path)
    if not candidate.is_absolute():
        candidate = (markdown_dir / candidate).resolve()
    return candidate


def _build_table_caption(chapter: int, number: int, title: str) -> str:
    return f"表{chapter}-{number} {_normalize_caption_text(title, '表格')}"


def _build_continued_table_label(chapter: int, number: int) -> str:
    return f"续表{chapter}－{number}"


def _build_figure_caption(chapter: int, number: int, title: str, image_path: Path) -> str:
    fallback = image_path.stem.replace("_", " ").strip() or "插图"
    return f"图{chapter}-{number} {_normalize_caption_text(title, fallback)}"


def _visual_length(text: str) -> int:
    length = 0
    for ch in text:
        if ch == "\n":
            continue
        if ch.isspace():
            length += 1
        elif ord(ch) < 128:
            length += 1
        else:
            length += 2
    return length


def _estimate_table_row_units(row: List[str], column_count: int) -> int:
    max_units = 1
    per_cell_capacity = max(10, int(30 / max(column_count, 1)))
    for cell in row:
        pieces = str(cell).splitlines() or [""]
        cell_units = 0
        for piece in pieces:
            cell_units += max(1, (_visual_length(piece) + per_cell_capacity - 1) // per_cell_capacity)
        max_units = max(max_units, cell_units)
    return max_units


def _split_table_rows(rows: List[List[str]]) -> List[List[List[str]]]:
    if not ENABLE_HEURISTIC_CONTINUED_TABLES:
        return [rows]

    if len(rows) <= 2:
        return [rows]

    header = rows[0]
    body_rows = rows[1:]
    column_count = len(header)
    chunks: List[List[List[str]]] = []
    current_body: List[List[str]] = []
    current_units = 0
    capacity = FIRST_TABLE_PAGE_CAPACITY

    for row in body_rows:
        row_units = _estimate_table_row_units(row, column_count)
        if current_body and current_units + row_units > capacity:
            chunks.append([header] + current_body)
            current_body = [row]
            current_units = row_units
            capacity = CONTINUED_TABLE_PAGE_CAPACITY
            continue
        current_body.append(row)
        current_units += row_units

    if current_body:
        chunks.append([header] + current_body)

    return chunks or [rows]


def _write_table_chunk(doc: Document, end_paragraph, rows: List[List[str]]) -> None:
    table = insert_table_before(doc, end_paragraph, len(rows), len(rows[0]), style="Normal Table")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])

    section = _body_section_for_media(doc)
    cell_width = Emu(int(_available_body_width(section).emu / max(len(rows[0]), 1)))

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.width = cell_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_text(cell, value, TABLE_CELL_STYLE)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            paragraph.paragraph_format.line_spacing = 1.0


def _write_table(doc: Document, end_paragraph, prototype, chapter: int, number: int, rows: List[List[str]], caption: str, caption_paragraph=None) -> None:
    if not rows or not rows[0]:
        return

    chunks = _split_table_rows(rows)
    for chunk_index, chunk_rows in enumerate(chunks):
        if chunk_index == 0:
            active_caption_paragraph = caption_paragraph
            if active_caption_paragraph is None:
                active_caption_paragraph = insert_paragraph_before(end_paragraph, prototype)
            set_paragraph_text(active_caption_paragraph, _build_table_caption(chapter, number, caption), TABLE_CAPTION_STYLE)
            format_table_caption(active_caption_paragraph)
        else:
            continued_label = insert_paragraph_before(end_paragraph, prototype)
            set_paragraph_text(continued_label, _build_continued_table_label(chapter, number), CONTINUED_TABLE_LABEL_STYLE)
            continued_label.paragraph_format.page_break_before = True
            format_continued_table_label(continued_label)
        _write_table_chunk(doc, end_paragraph, chunk_rows)


def _write_image(doc: Document, end_paragraph, prototype, chapter: int, number: int, block, markdown_dir: Path, picture_paragraph=None) -> None:
    image_path = _resolve_media_path(block.path, markdown_dir)
    if not image_path.exists():
        raise FileNotFoundError(f"图片文件不存在: {image_path}")

    if picture_paragraph is None:
        picture_paragraph = insert_paragraph_before(end_paragraph, prototype)
    clear_paragraph(picture_paragraph)
    section = _body_section_for_media(doc)
    max_width = Emu(_available_body_width(section).emu - Cm(1).emu)
    preferred_width = Cm(11.5)
    picture_width = preferred_width if preferred_width.emu < max_width.emu else max_width
    picture_paragraph.add_run().add_picture(str(image_path), width=picture_width)
    format_picture_block(picture_paragraph)

    caption_paragraph = insert_paragraph_before(end_paragraph, prototype)
    set_paragraph_text(
        caption_paragraph,
        _build_figure_caption(chapter, number, block.caption, image_path),
        FIGURE_CAPTION_STYLE,
    )
    format_figure_caption(caption_paragraph)


def _write_body(doc: Document, anchors: Dict[str, int], body_blocks, markdown_dir: Path) -> None:
    reference_paragraph = doc.paragraphs[anchors["reference"]]
    prototypes = {
        "h1": doc.paragraphs[anchors["body_start"]],
        "h2": doc.paragraphs[anchors["body_start"] + 1] if anchors["body_start"] + 1 < anchors["reference"] else doc.paragraphs[anchors["body_start"]],
        "h3": doc.paragraphs[anchors["body_start"] + 2] if anchors["body_start"] + 2 < anchors["reference"] else doc.paragraphs[anchors["body_start"]],
        "body": doc.paragraphs[anchors["body_start"] + 3] if anchors["body_start"] + 3 < anchors["reference"] else doc.paragraphs[anchors["body_start"]],
    }
    end_paragraph = reference_paragraph
    start_paragraph = _prepare_slot(doc, doc.paragraphs[anchors["body_start"]], reference_paragraph, preserve_start=True)

    chapter = 0
    section = 0
    subsection = 0
    table_number = 0
    figure_number = 0
    current_paragraph = start_paragraph
    first_written = False

    def next_target(prototype_key: str):
        nonlocal current_paragraph, first_written
        prototype = prototypes[prototype_key]
        if not first_written:
            copy_paragraph_properties(current_paragraph, prototype)
            clear_paragraph(current_paragraph)
            first_written = True
            return current_paragraph
        current_paragraph = insert_paragraph_before(end_paragraph, prototype)
        return current_paragraph

    for block in body_blocks:
        if block.type == "heading":
            text = block.text
            if block.level == 1:
                chapter += 1
                section = 0
                subsection = 0
                table_number = 0
                figure_number = 0
                numbered = f"{chapter}  {text}"
                paragraph = next_target("h1")
                set_paragraph_text(paragraph, numbered, MAJOR_HEADING_STYLE)
                format_major_heading(paragraph)
            elif block.level == 2:
                section += 1
                subsection = 0
                numbered = f"{chapter}.{section} {text}"
                paragraph = next_target("h2")
                set_paragraph_text(paragraph, numbered, SUBHEADING_STYLE)
                format_subheading(paragraph)
            else:
                subsection += 1
                numbered = f"{chapter}.{section}.{subsection} {text}"
                paragraph = next_target("h3")
                set_paragraph_text(paragraph, numbered, SUBHEADING_STYLE)
                format_subheading(paragraph)
            continue

        active_chapter = chapter or 1
        if block.type == "table":
            table_number += 1
            caption_paragraph = None if first_written else next_target("body")
            _write_table(
                doc,
                end_paragraph,
                prototypes["body"],
                active_chapter,
                table_number,
                block.rows,
                block.caption,
                caption_paragraph=caption_paragraph,
            )
            continue

        if block.type == "image":
            figure_number += 1
            picture_paragraph = None if first_written else next_target("body")
            _write_image(
                doc,
                end_paragraph,
                prototypes["body"],
                active_chapter,
                figure_number,
                block,
                markdown_dir,
                picture_paragraph=picture_paragraph,
            )
            continue

        paragraph = next_target("body")
        set_paragraph_text(paragraph, block.text, BODY_STYLE, superscript_citations=True)
        format_body(paragraph)

    if not body_blocks:
        clear_paragraph(start_paragraph)


def fill_document(template_path: Path, markdown_path: Path, output_path: Path, schema_path: Path, metadata_path: Optional[Path]) -> None:
    schema = load_schema(schema_path)
    paper = parse_markdown(markdown_path)
    metadata = load_metadata(metadata_path)
    paper = merge_content_with_metadata(paper, metadata)

    doc = Document(str(template_path))
    anchors = locate_anchors(doc, schema, required=True)

    fill_cover_fields(doc, schema, paper.cover, blank_missing=True)
    english_title_lines = wrap_english_title(paper.english_title) if paper.english_title else ["", "", ""]
    set_cover_titles(
        doc,
        anchors["cn_abstract"],
        anchors["cn_keywords"],
        anchors["en_abstract"],
        paper.title,
        english_title_lines,
    )
    update_headers(doc, paper.title)

    anchors = locate_anchors(doc, schema, required=True)
    cn_end = doc.paragraphs[anchors["cn_keywords"]]
    cn_start = _prepare_slot(doc, doc.paragraphs[anchors["cn_abstract"]], cn_end, preserve_start=False)
    _write_text_blocks(cn_start, paper.cn_abstract, BODY_STYLE, cn_end, formatter=format_body)
    set_keyword_paragraph(cn_end, "关键词：", "，".join(paper.cn_keywords), KEYWORD_LABEL_STYLE, BODY_STYLE)
    format_keywords(cn_end)

    anchors = locate_anchors(doc, schema, required=True)
    en_end = doc.paragraphs[anchors["en_keywords"]]
    en_start = _prepare_slot(doc, doc.paragraphs[anchors["en_abstract"]], en_end, preserve_start=False)
    _write_text_blocks(en_start, paper.en_abstract, ENGLISH_BODY_STYLE, en_end, formatter=format_body)
    set_keyword_paragraph(en_end, "KEY WORDS: ", ", ".join(paper.en_keywords), ENGLISH_KEYWORD_LABEL_STYLE, ENGLISH_BODY_STYLE)
    format_keywords(en_end)

    anchors = locate_anchors(doc, schema, required=True)
    toc_start = _prepare_slot(doc, doc.paragraphs[anchors["toc"]], doc.paragraphs[anchors["body_start"]], preserve_start=False)
    clear_paragraph(toc_start)

    anchors = locate_anchors(doc, schema, required=True)
    reference_end_name = find_next_anchor(anchors, "reference", "appendix", "acknowledgements", "foreign_translation")
    reference_end_idx = anchors[reference_end_name] if reference_end_name else len(doc.paragraphs) - 1
    format_major_heading_on_new_page(doc.paragraphs[anchors["reference"]])
    ref_end = doc.paragraphs[reference_end_idx]
    ref_start = _prepare_slot(doc, doc.paragraphs[anchors["reference"]], ref_end, preserve_start=False)
    _write_references(doc, ref_start, paper.references, ref_end)

    optional_sections = [
        ("foreign_translation", "foreign_translation_body"),
        ("acknowledgements", "acknowledgements_body", "foreign_translation"),
        ("appendix", "appendix_body", "acknowledgements", "foreign_translation"),
    ]
    for section_name, _, *next_candidates in optional_sections:
        anchors = locate_anchors(doc, schema, required=True)
        start_idx = anchors.get(section_name)
        if start_idx is None:
            continue
        end_name = find_next_anchor(anchors, section_name, *next_candidates)
        end_idx = anchors[end_name] if end_name else len(doc.paragraphs) - 1
        section_end = doc.paragraphs[end_idx]
        section_start = _prepare_slot(doc, doc.paragraphs[start_idx], section_end, preserve_start=False)
        paragraphs = getattr(paper, section_name)
        _write_text_blocks(section_start, paragraphs, BODY_STYLE, section_end, formatter=format_body)

    anchors = locate_anchors(doc, schema, required=True)
    _write_body(doc, anchors, paper.body, markdown_path.parent)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="基于东华本科论文母版进行内容直填")
    parser.add_argument("markdown", help="输入 Markdown 路径")
    parser.add_argument("-t", "--template", default=str(DEFAULT_TEMPLATE), help="干净母版路径")
    parser.add_argument("-o", "--output", help="输出 docx 路径")
    parser.add_argument("-s", "--schema", default=str(DEFAULT_SCHEMA), help="模板 schema 路径")
    parser.add_argument("-m", "--meta", help="元数据 JSON 路径")
    args = parser.parse_args()

    markdown_path = Path(args.markdown).resolve()
    template_path = Path(args.template).resolve()
    schema_path = Path(args.schema).resolve()
    metadata_path = Path(args.meta).resolve() if args.meta else None
    output_path = Path(args.output).resolve() if args.output else SCRIPT_DIR / "outputs" / f"{markdown_path.stem}_东华直填.docx"

    fill_document(template_path, markdown_path, output_path, schema_path, metadata_path)
    print(f"已生成: {output_path}")


if __name__ == "__main__":
    main()
