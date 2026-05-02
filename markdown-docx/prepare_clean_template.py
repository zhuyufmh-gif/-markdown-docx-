#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
from pathlib import Path

from docx import Document

from cover_field_locator import fill_cover_fields, set_cover_titles
from docx_utils import (
    RunStyle,
    body_blocks_between,
    clear_paragraph,
    first_nonempty_run_style,
    remove_or_clear_block,
    set_keyword_paragraph,
    set_paragraph_text,
)
from fill_template import _prepare_slot
from header_utils import update_headers
from template_locator import find_next_anchor, load_schema, locate_anchors


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_SCHEMA = SCRIPT_DIR / "dhu_template_schema.json"
DEFAULT_SOURCE = SCRIPT_DIR.parent / "本科毕业论文参考模版（2025版）.docx"
DEFAULT_OUTPUT = SCRIPT_DIR / "outputs" / "dhu_undergrad_clean_template.docx"


def build_clean_template(source_path: Path, output_path: Path, schema_path: Path) -> None:
    schema = load_schema(schema_path)
    placeholders = schema["clean_placeholders"]
    doc = Document(str(source_path))
    anchors = locate_anchors(doc, schema, required=True)

    fill_cover_fields(
        doc,
        schema,
        {
            "college": "",
            "major": "",
            "author": "",
            "student_id": "",
            "supervisor": "",
            "submit_date": "",
        },
        blank_missing=True,
    )
    set_cover_titles(
        doc,
        anchors["cn_abstract"],
        anchors["cn_keywords"],
        anchors["en_abstract"],
        placeholders["cover_cn_title"],
        placeholders["front_en_title_lines"],
    )
    update_headers(doc, placeholders["front_cn_title"])

    anchors = locate_anchors(doc, schema, required=True)
    cn_keywords = doc.paragraphs[anchors["cn_keywords"]]
    cn_start = _prepare_slot(doc, doc.paragraphs[anchors["cn_abstract"]], cn_keywords, preserve_start=False)
    cn_style = first_nonempty_run_style(cn_start, fallback_font="宋体", fallback_size=12.0)
    set_paragraph_text(cn_start, placeholders["cn_abstract_body"], cn_style)
    cn_prefix_style = first_nonempty_run_style(cn_keywords, fallback_font="黑体", fallback_size=12.0)
    set_keyword_paragraph(cn_keywords, placeholders["cn_keywords"], "", cn_prefix_style, RunStyle(size_pt=12.0, bold=False))

    anchors = locate_anchors(doc, schema, required=True)
    en_keywords = doc.paragraphs[anchors["en_keywords"]]
    en_start = _prepare_slot(doc, doc.paragraphs[anchors["en_abstract"]], en_keywords, preserve_start=False)
    en_style = first_nonempty_run_style(en_start, fallback_font="Times New Roman", fallback_size=12.0)
    set_paragraph_text(en_start, placeholders["en_abstract_body"], en_style)
    en_prefix_style = first_nonempty_run_style(en_keywords, fallback_font="Times New Roman", fallback_size=12.0)
    set_keyword_paragraph(en_keywords, placeholders["en_keywords"], "", en_prefix_style, RunStyle(ascii_font="Times New Roman", east_asia_font="Times New Roman", size_pt=12.0, bold=False))

    anchors = locate_anchors(doc, schema, required=True)
    toc_slot = _prepare_slot(doc, doc.paragraphs[anchors["toc"]], doc.paragraphs[anchors["body_start"]], preserve_start=False)
    clear_paragraph(toc_slot)

    anchors = locate_anchors(doc, schema, required=True)
    body_start_paragraph = doc.paragraphs[anchors["body_start"]]
    body_h2 = doc.paragraphs[anchors["body_start"] + 1]
    body_h3 = doc.paragraphs[anchors["body_start"] + 2]
    body_text = doc.paragraphs[anchors["body_start"] + 3]
    for block in body_blocks_between(doc, body_text, doc.paragraphs[anchors["reference"]]):
        remove_or_clear_block(block)

    body_h1_style = first_nonempty_run_style(body_start_paragraph, fallback_font="黑体", fallback_size=16.0)
    set_paragraph_text(body_start_paragraph, placeholders["body_h1"], body_h1_style)

    body_h2_style = first_nonempty_run_style(body_h2, fallback_font="黑体", fallback_size=14.0)
    set_paragraph_text(body_h2, placeholders["body_h2"], body_h2_style)
    body_h3_style = first_nonempty_run_style(body_h3, fallback_font="黑体", fallback_size=14.0)
    set_paragraph_text(body_h3, placeholders["body_h3"], body_h3_style)
    body_text_style = first_nonempty_run_style(body_text, fallback_font="宋体", fallback_size=12.0)
    set_paragraph_text(body_text, placeholders["body_text"], body_text_style)

    anchors = locate_anchors(doc, schema, required=True)
    reference_end_name = find_next_anchor(anchors, "reference", "appendix", "acknowledgements", "foreign_translation")
    reference_end_idx = anchors[reference_end_name] if reference_end_name else len(doc.paragraphs) - 1
    reference_end = doc.paragraphs[reference_end_idx]
    ref_slot = _prepare_slot(doc, doc.paragraphs[anchors["reference"]], reference_end, preserve_start=False)
    ref_style = first_nonempty_run_style(ref_slot, fallback_font="宋体", fallback_size=12.0)
    set_paragraph_text(ref_slot, placeholders["reference_item"], ref_style)

    for section_name, placeholder_key, *next_candidates in [
        ("foreign_translation", "foreign_translation_body"),
        ("acknowledgements", "acknowledgements_body", "foreign_translation"),
        ("appendix", "appendix_body", "acknowledgements", "foreign_translation"),
    ]:
        anchors = locate_anchors(doc, schema, required=True)
        start_idx = anchors.get(section_name)
        if start_idx is None:
            continue
        end_name = find_next_anchor(anchors, section_name, *next_candidates)
        end_idx = anchors[end_name] if end_name else len(doc.paragraphs) - 1
        section_end = doc.paragraphs[end_idx]
        slot = _prepare_slot(doc, doc.paragraphs[start_idx], section_end, preserve_start=False)
        style = first_nonempty_run_style(slot, fallback_font="宋体", fallback_size=12.0)
        set_paragraph_text(slot, placeholders[placeholder_key], style)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="从东华原始模板生成可长期复用的干净母版")
    parser.add_argument("-i", "--input", default=str(DEFAULT_SOURCE), help="原始模板路径")
    parser.add_argument("-o", "--output", default=str(DEFAULT_OUTPUT), help="干净母版输出路径")
    parser.add_argument("-s", "--schema", default=str(DEFAULT_SCHEMA), help="schema 路径")
    args = parser.parse_args()

    build_clean_template(Path(args.input).resolve(), Path(args.output).resolve(), Path(args.schema).resolve())
    print(f"已生成干净母版: {Path(args.output).resolve()}")


if __name__ == "__main__":
    main()
