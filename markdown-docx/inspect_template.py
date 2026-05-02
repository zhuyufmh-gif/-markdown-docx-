#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
from pathlib import Path

from docx import Document

from template_locator import load_schema, locate_anchors


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_SCHEMA = SCRIPT_DIR / "dhu_template_schema.json"


def main() -> None:
    parser = argparse.ArgumentParser(description="输出模板锚点与基础结构摘要")
    parser.add_argument("template", help="待分析 docx 路径")
    parser.add_argument("-s", "--schema", default=str(DEFAULT_SCHEMA), help="schema 路径")
    args = parser.parse_args()

    doc = Document(str(Path(args.template).resolve()))
    schema = load_schema(Path(args.schema).resolve())
    anchors = locate_anchors(doc, schema, required=False)

    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"sections={len(doc.sections)}")
    for name, index in sorted(anchors.items(), key=lambda item: item[1]):
        print(f"{name}: {index} -> {doc.paragraphs[index].text}")


if __name__ == "__main__":
    main()
